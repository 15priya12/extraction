from docx import Document
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
import os
import sys
import re
from collections import defaultdict
import xml.etree.ElementTree as ET

sys.dont_write_bytecode = True
os.environ['PYTHONDONTWRITEBYTECODE'] = '1'


def to_roman(num, lowercase=False):
    """Convert number to Roman numeral."""
    vals = [
        (1000, "M"), (900, "CM"), (500, "D"), (400, "CD"),
        (100, "C"), (90, "XC"), (50, "L"), (40, "XL"),
        (10, "X"), (9, "IX"), (5, "V"), (4, "IV"), (1, "I")
    ]
    result = ""
    for v, sym in vals:
        while num >= v:
            result += sym
            num -= v
    return result.lower() if lowercase else result


def to_letter(num, lowercase=False):
    """Convert number to letter sequence (A, B, C, ..., Z, AA, AB, etc.)."""
    result = ""
    while num > 0:
        num, rem = divmod(num - 1, 26)
        result = chr(65 + rem) + result
    return result.lower() if lowercase else result


class GenerateParaRefsForDocx:
    def __init__(self, start_index: int = 100000):
        self.data = []
        self.tables = []
        self.table_counter = 0
        self.para_id_counter = start_index
        self.namespaces = {
            'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
            'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
        }
        # Numbering related attributes
        self.numbering_map = defaultdict(dict)
        self.global_counters = defaultdict(int)  # Track counters globally by level
        self.global_stack = {}  # Track the current hierarchy globally
        self.last_numId = None
        self.last_ilvl = -1

    def get_numbering(self, doc: Document) -> defaultdict:
        """Extract numbering information from the document."""
        numbering_map = defaultdict(dict)
        try:
            numbering_part = doc.part.numbering_part
            if not numbering_part:
                return numbering_map

            root = ET.fromstring(numbering_part.blob)
            abstract_nums = {}

            for absnum in root.findall("w:abstractNum", self.namespaces):
                abs_id = absnum.attrib.get(f"{{{self.namespaces['w']}}}abstractNumId")
                abs_levels = {}
                for lvl in absnum.findall("w:lvl", self.namespaces):
                    ilvl = lvl.attrib.get(f"{{{self.namespaces['w']}}}ilvl")
                    num_fmt_el = lvl.find("w:numFmt", self.namespaces)
                    lvl_text_el = lvl.find("w:lvlText", self.namespaces)
                    fmt = num_fmt_el.attrib.get(f"{{{self.namespaces['w']}}}val") if num_fmt_el is not None else "bullet"
                    pattern = lvl_text_el.attrib.get(f"{{{self.namespaces['w']}}}val") if lvl_text_el is not None else "•"
                    abs_levels[ilvl] = (fmt, pattern)
                abstract_nums[abs_id] = abs_levels

            for num in root.findall("w:num", self.namespaces):
                num_id = num.attrib.get(f"{{{self.namespaces['w']}}}numId")
                abs_id_el = num.find("w:abstractNumId", self.namespaces)
                if abs_id_el is not None:
                    abs_id = abs_id_el.attrib.get(f"{{{self.namespaces['w']}}}val")
                    numbering_map[num_id] = abstract_nums.get(abs_id, {})

        except Exception as e:
            print(f"Error extracting numbering information: {e}")

        return numbering_map

    def generate_guid(self) -> str:
        """Generate an incremental ID for para_id starting from the configured start_index"""
        current_id = self.para_id_counter
        self.para_id_counter += 1
        return str(current_id)

    def extract_insertions_only(self, paragraph: Paragraph) -> str:
        """Extract paragraph text preserving insertions, skipping deletions,
        and keeping only the visible hyperlink text (not the target URL)."""
        parts = []

        for child in paragraph._element:
            tag = child.tag

            if tag.endswith('}hyperlink'):
                text_runs = child.findall('.//w:t', self.namespaces)
                display_text = "".join([t.text for t in text_runs if t.text])
                if display_text:
                    parts.append(display_text)

            elif tag.endswith('}r'):
                for t in child.findall('.//w:t', self.namespaces):
                    if t.text:
                        parts.append(t.text)

            elif tag.endswith('}ins'):
                for t in child.findall('.//w:t', self.namespaces):
                    if t.text:
                        parts.append(t.text)

            elif tag.endswith('}del'):
                continue

        return "".join(parts).strip()

    def get_bullet_number(self, paragraph: Paragraph) -> str:
        """Extract bullet or number from a paragraph."""
        numPr = paragraph._element.find(".//w:numPr", self.namespaces)
        if numPr is None:
            return ""

        numId_el = numPr.find("w:numId", self.namespaces)
        ilvl_el = numPr.find("w:ilvl", self.namespaces)
        numId = numId_el.attrib.get(f"{{{self.namespaces['w']}}}val") if numId_el is not None else None
        ilvl = ilvl_el.attrib.get(f"{{{self.namespaces['w']}}}val") if ilvl_el is not None else "0"
        ilvl_int = int(ilvl)

        if not numId:
            return ""

        # Check if format changed from previous
        fmt_level, _ = self.numbering_map.get(numId, {}).get(str(ilvl_int), ("decimal", "%1"))
        prev_fmt = self.numbering_map.get(self.last_numId, {}).get(str(ilvl_int), ("decimal", "%1"))[0] if self.last_numId else None

        # Reset counter if format changed at same level
        if self.last_numId and self.last_ilvl == ilvl_int and fmt_level != prev_fmt:
            self.global_counters[ilvl_int] = 0

        # Increment current level
        self.global_counters[ilvl_int] += 1

        # Reset deeper levels
        for deeper in range(ilvl_int + 1, 9):
            self.global_counters[deeper] = 0
            self.global_stack.pop(deeper, None)

        # Build label parts for all levels from 0 to current
        label_parts = []
        for level in range(ilvl_int + 1):
            n = self.global_counters[level]

            # Get format from the current numId
            fmt_level, _ = self.numbering_map.get(numId, {}).get(str(level), ("decimal", "%1"))

            # Get format from previous level if it exists
            prev_level_fmt = None
            if level > 0:
                prev_level_fmt = self.numbering_map.get(numId, {}).get(str(level - 1), ("decimal", "%1"))[0]

            # Handle numbering format
            if fmt_level == "decimal":
                val = str(n)
            elif fmt_level == "upperLetter":
                val = to_letter(n)
            elif fmt_level == "lowerLetter":
                # Only use letter if previous level was also letter or if explicitly configured
                if prev_level_fmt and prev_level_fmt.endswith("Letter"):
                    val = to_letter(n, lowercase=True)
                else:
                    val = str(n)
            elif fmt_level == "upperRoman":
                val = to_roman(n)
            elif fmt_level == "lowerRoman":
                val = to_roman(n, lowercase=True)
            else:
                val = str(n)

            label_parts.append(val)
            self.global_stack[level] = val

        self.last_numId = numId
        self.last_ilvl = ilvl_int

        return ".".join(label_parts)

    def extract_text_from_paragraph(self, paragraph: Paragraph) -> tuple[str, str]:
        """Extract text and bullet/number from paragraph."""
        bullet = self.get_bullet_number(paragraph)
        text = self.extract_insertions_only(paragraph)
        
        # Remove the numbering from the text if it's already there
        if bullet:
            text = re.sub(r'^[\d\w]+(?:\.[\d\w]+)*[\.\)]\s*', '', text)
            
        return bullet, text

    def extract_cell_content(self, cell: _Cell) -> str:
        """Extract content from a table cell, handling nested tables and paragraphs"""
        content_parts = []
        for element in cell._element:
            if element.tag.endswith('}p'):
                paragraph = Paragraph(element, cell)
                bullet, para_text = self.extract_text_from_paragraph(paragraph)
                if para_text:
                    if bullet:
                        content_parts.append(f"{bullet} {para_text}")
                    else:
                        content_parts.append(para_text)
            elif element.tag.endswith('}tbl'):
                nested_table = Table(element, cell)
                nested_table_md = self.process_table(nested_table)
                if nested_table_md:
                    content_parts.append(f"\n{nested_table_md}\n")

        return " ".join(content_parts).strip()

    def process_table(self, table: Table) -> str:
        """Process a table and convert it to markdown table format, skipping empty tables."""
        markdown_table = []
        rows = table.rows
        if not rows:
            return ""

        header_cells = [self.extract_cell_content(cell) for cell in rows[0].cells]
        if not any(cell.strip() for cell in header_cells):
            return ""

        markdown_table.append("| " + " | ".join(header_cells) + " |")
        markdown_table.append("|" + " --- |" * len(header_cells))

        data_rows_added = False
        for row in rows[1:]:
            row_cells = [self.extract_cell_content(cell) for cell in row.cells]
            if any(cell.strip() for cell in row_cells):
                markdown_table.append("| " + " | ".join(row_cells) + " |")
                data_rows_added = True

        if not data_rows_added and all(not c.strip() for c in header_cells):
            return ""

        return "\n".join(markdown_table) if markdown_table else ""

    def extract_plain_text(self, file_path: str) -> str:
        """Fallback method to extract plain text from DOCX"""
        try:
            doc = Document(file_path)
            full_text = []
            table_index = 0

            # Initialize numbering for plain text extraction
            self.numbering_map = self.get_numbering(doc)
            self.global_counters.clear()
            self.global_stack.clear()
            self.last_numId = None
            self.last_ilvl = -1

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    bullet, para_text = self.extract_text_from_paragraph(paragraph)
                    if para_text:
                        if bullet:
                            full_text.append(f"{bullet} {para_text}")
                        else:
                            full_text.append(para_text)

            for table in doc.tables:
                table_index += 1
                table_markdown = self.process_table(table)
                if table_markdown:
                    self.tables.append({
                        'index': table_index,
                        'content': table_markdown
                    })
                    full_text.append(f"SEE TABLE_#{table_index} below")

            result = "\n\n".join(full_text)

            if self.tables:
                result += "\n\n"
                for table_info in self.tables:
                    result += f"\nTABLE_#{table_info['index']}:\n{table_info['content']}\n"

            return result

        except Exception as e:
            print(f"Error in extract_plain_text: {e}")
            return ""

    def process_docx_file(self, file_path: str, start_index: int) -> str:
        """Process a DOCX file and return markdown content including headers and short-title detection"""
        try:
            self.para_id_counter = start_index
            doc = Document(file_path)
            
            # Initialize numbering
            self.numbering_map = self.get_numbering(doc)
            self.global_counters.clear()
            self.global_stack.clear()
            self.last_numId = None
            self.last_ilvl = -1
            
            current_header = ""

            for element in doc.element.body:
                if element.tag.endswith('}p'):
                    paragraph = Paragraph(element, doc)
                    bullet, para_text = self.extract_text_from_paragraph(paragraph)
                    if not para_text:
                        continue

                    style_name = getattr(paragraph.style, 'name', '').lower() if paragraph.style else ''
                    word_count = len(para_text.split())

                    # Check if it's a top-level bullet (no dots in the bullet)
                    is_top_level = bullet == "" or "." not in bullet

                    # Update header if:
                    # 1. It's a heading style, OR
                    # 2. It's a short text (≤6 words) AND either has no bullet or has a top-level bullet
                    if "heading" in style_name or (word_count <= 6 and is_top_level):
                        current_header = para_text.strip()
                    # If it's a sub-item (has dots in bullet), don't update the header
                    elif not is_top_level:
                        pass  # Keep the current header

                    # Add paragraph regardless of whether it's a header
                    self.data.append({
                        'para_id': self.generate_guid(),
                        'header': current_header,
                        'para_text': para_text,
                        'bullet': bullet
                    })

                elif element.tag.endswith('}tbl'):
                    self.table_counter += 1
                    table = Table(element, doc)
                    table_markdown = self.process_table(table)

                    if table_markdown:
                        self.tables.append({
                            'index': self.table_counter,
                            'content': table_markdown
                        })
                        self.data.append({
                            'para_id': self.generate_guid(),
                            'header': current_header,
                            'para_text': f"SEE TABLE_{self.table_counter} below"
                        })

            if self.data:
                return self.generate_markdown_table_with_header()
            else:
                return self.extract_plain_text(file_path)

        except Exception as e:
            print(f"Error processing DOCX: {e}")
            return self.extract_plain_text(file_path)

    def generate_markdown_table_with_header(self) -> str:
        """Generate markdown table with headers"""
        if not self.data:
            return ""

        lines = ["| para_id | header | bullet | para_text |", "|---------|---------|---------|-----------|"]

        for item in self.data:
            header = item.get('header', '')
            bullet = item.get('bullet', '')
            escaped_text = item['para_text'].replace('|', '\\|').replace('\n', '<br>')
            escaped_header = header.replace('|', '\\|').replace('\n', '<br>')
            escaped_bullet = bullet.replace('|', '\\|')
            lines.append(f"| {item['para_id']} | {escaped_header} | {escaped_bullet} | {escaped_text} |")

        result = "\n".join(lines)

        if self.tables:
            result += "\n\n"
            for table_info in self.tables:
                result += f"\nTABLE_{table_info['index']}:\n{table_info['content']}\n"

        return result

    def clear_data(self):
        """Clear the stored data"""
        self.data = []
        self.tables = []
        self.table_counter = 0


if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser(description="Extract DOCX content with paragraph references and inferred headers.")
    parser.add_argument("input_file", help="Path to the DOCX file to process.")
    parser.add_argument("--start_index", type=int, default=100000, help="Starting index for paragraph IDs.")
    args = parser.parse_args()

    processor = GenerateParaRefsForDocx(start_index=args.start_index)
    markdown_output = processor.process_docx_file(args.input_file, args.start_index)

    output_file = os.path.splitext(args.input_file)[0] + "_output.md"
    print("sucess")
    with open(output_file, "w", encoding="utf-8") as f:
        f.write(markdown_output)